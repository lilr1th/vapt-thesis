#!/usr/bin/env python3
"""
generate_docx.py — VAPT Thesis → ITC-compliant Word Document
ITC Format: Left 3cm, Right 2cm, Top 2cm, Bottom 2cm
Headings: Chapter=16pt bold ALL CAPS | X.X=14pt bold | X.X.X=12pt bold
Body: 12pt TNR, 1.5 spacing, 6pt after, justified

Requirements: pip install python-docx
After opening .docx in Word: press Ctrl+A then F9 to update Table of Contents.
"""
import os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE  = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(BASE, "screenshots")
IMGS  = os.path.join(BASE, "image")
MD    = os.path.join(BASE, "THESIS_VAPT_FULL.md")
OUT   = os.path.join(BASE, "THESIS_VAPT_WORD.docx")

# ── Color palette ─────────────────────────────────────────────────────────────
C = {
    "critical": ("8B0000", "FFFFFF"),
    "high":     ("C00000", "FFFFFF"),
    "medium":   ("FFC000", "000000"),
    "low":      ("70AD47", "FFFFFF"),
    "label":    ("D9D9D9", "000000"),
    "blue":     ("2E74B5", "FFFFFF"),
}
def hex2rgb(h): return RGBColor(int(h[:2],16), int(h[2:4],16), int(h[4:],16))

# ── XML helpers ───────────────────────────────────────────────────────────────
def set_bg(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')):
        tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_vert(cell, val='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    v = OxmlElement('w:vAlign')
    v.set(qn('w:val'), val)
    tcPr.append(v)

def add_section_break(doc, restart_arabic=False):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    for tag, val in [
        ('w:pgSz',  {'w:w': str(int(21.0/2.54*1440)), 'w:h': str(int(29.7/2.54*1440))}),
        ('w:pgMar', {'w:top': str(int(2.0/2.54*1440)), 'w:right': str(int(2.0/2.54*1440)),
                     'w:bottom': str(int(2.0/2.54*1440)), 'w:left': str(int(3.0/2.54*1440))}),
    ]:
        el = OxmlElement(tag)
        for k, v2 in val.items():
            el.set(qn(k), v2)
        sectPr.append(el)
    if restart_arabic:
        pg = OxmlElement('w:pgNumType')
        pg.set(qn('w:fmt'), 'decimal')
        pg.set(qn('w:start'), '1')
        sectPr.append(pg)
    pPr.append(sectPr)

def set_toc_styles(doc):
    """Inject TOC1/TOC2/TOC3 style XML directly so Word recognises them on F9."""
    # styleId must be 'TOC1' (no space) — python-docx add_style() uses wrong id
    # Tab at 9072 twips = 16cm = full text width (A4 minus 3cm left + 2cm right)
    styles_el = doc.part.styles._element

    # Remove any previously injected TOC styles to avoid duplicates
    for st in list(styles_el):
        nm = st.find(qn('w:name'))
        if nm is not None and nm.get(qn('w:val'), '').lower() in ('toc 1','toc 2','toc 3'):
            styles_el.remove(st)

    configs = [
        # (styleId, display name, left indent twips, bold, all-caps, font size pt)
        ('TOC1', 'toc 1',    0,   True,  True,  16),
        ('TOC2', 'toc 2',  284,   True,  False, 14),
        ('TOC3', 'toc 3',  851,   False, False, 12),
    ]
    for sid, sname, indent, bold, caps, size in configs:
        st = OxmlElement('w:style')
        st.set(qn('w:type'),    'paragraph')
        st.set(qn('w:styleId'), sid)

        nm = OxmlElement('w:name'); nm.set(qn('w:val'), sname); st.append(nm)
        bo = OxmlElement('w:basedOn'); bo.set(qn('w:val'), 'Normal'); st.append(bo)

        pPr = OxmlElement('w:pPr')
        sp  = OxmlElement('w:spacing')
        sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '60'); pPr.append(sp)
        if indent:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(indent)); ind.set(qn('w:firstLine'), '0')
            pPr.append(ind)
        tabs = OxmlElement('w:tabs')
        tab  = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right'); tab.set(qn('w:leader'), 'dot')
        tab.set(qn('w:pos'), '9072')
        tabs.append(tab); pPr.append(tabs)
        st.append(pPr)

        sz_val = str(size * 2)  # half-points
        rPr   = OxmlElement('w:rPr')
        fonts = OxmlElement('w:rFonts')
        fonts.set(qn('w:ascii'), 'Times New Roman')
        fonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(fonts)
        if bold:
            rPr.append(OxmlElement('w:b'))
            rPr.append(OxmlElement('w:bCs'))
        if caps:
            rPr.append(OxmlElement('w:caps'))
        sz = OxmlElement('w:sz');   sz.set(qn('w:val'),   sz_val); rPr.append(sz)
        sc = OxmlElement('w:szCs'); sc.set(qn('w:val'),   sz_val); rPr.append(sc)
        st.append(rPr)

        styles_el.append(st)

def add_toc_field(doc):
    set_toc_styles(doc)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin'); run._r.append(fc1)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    it.text = ' TOC \\o "1-3" \\h \\z \\u '; run._r.append(it)
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'separate'); run._r.append(fc2)
    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = Cm(0)
    r2 = p2.add_run('[Right-click → Update Field to generate Table of Contents]')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11)
    r2.font.italic = True; r2.font.color.rgb = RGBColor(0x80,0x80,0x80)
    run3 = p2.add_run()
    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'end'); run3._r.append(fc3)

# ── Document setup ────────────────────────────────────────────────────────────
def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    _setup_styles(doc)
    return doc

def _setup_styles(doc):
    n = doc.styles['Normal']
    n.font.name = 'Times New Roman'; n.font.size = Pt(12)
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    n.paragraph_format.space_after = Pt(6)
    defs = [
        ('Heading 1', 16, WD_ALIGN_PARAGRAPH.CENTER, True,  Pt(24), Pt(12)),
        ('Heading 2', 14, WD_ALIGN_PARAGRAPH.LEFT,   False, Pt(12), Pt(6)),
        ('Heading 3', 12, WD_ALIGN_PARAGRAPH.LEFT,   False, Pt(8),  Pt(4)),
        ('Heading 4', 12, WD_ALIGN_PARAGRAPH.LEFT,   False, Pt(6),  Pt(3)),
    ]
    for name, size, align, pb, sb, sa in defs:
        try:
            s = doc.styles[name]
            s.font.name = 'Times New Roman'; s.font.size = Pt(size)
            s.font.bold = True; s.font.color.rgb = RGBColor(0,0,0)
            s.paragraph_format.alignment = align
            s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            s.paragraph_format.space_before = sb; s.paragraph_format.space_after = sa
            s.paragraph_format.first_line_indent = Cm(0)
            if pb: s.paragraph_format.page_break_before = True
        except Exception: pass

# ── Typography helpers ────────────────────────────────────────────────────────
def fmt(run, size=12, bold=False, italic=False, color=None):
    run.font.name = 'Times New Roman'; run.font.size = Pt(size)
    run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = hex2rgb(color)

# Inline markdown → properly formatted runs (**bold**, *italic*, `code`, [link](url))
_INLINE_MD = re.compile(r'\*\*([^*]+)\*\*|\*([^*\[]+)\*|`([^`]+)`|\[([^\]]+)\]\([^)]+\)')
def add_inline_md(para, text, size=12):
    last = 0
    for m in _INLINE_MD.finditer(text):
        if m.start() > last:
            r = para.add_run(text[last:m.start()])
            r.font.name = 'Times New Roman'; r.font.size = Pt(size)
        if m.group(1):   # **bold**
            r = para.add_run(m.group(1))
            r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.bold = True
        elif m.group(2): # *italic*
            r = para.add_run(m.group(2))
            r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.italic = True
        elif m.group(3): # `code`
            r = para.add_run(m.group(3))
            r.font.name = 'Courier New'; r.font.size = Pt(size - 1)
        elif m.group(4): # [link text](url)
            r = para.add_run(m.group(4))
            r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.underline = True
        last = m.end()
    if last < len(text):
        r = para.add_run(text[last:])
        r.font.name = 'Times New Roman'; r.font.size = Pt(size)

def body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=1.27, size=12):
    if not text or not text.strip(): return
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    add_inline_md(p, text.strip(), size=size)
    return p

def ch_head(doc, text):
    """Chapter heading — 16pt bold ALL CAPS (body chapters only)."""
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    fmt(run, size=16, bold=True)

def front_head(doc, text):
    """Front matter heading — 12pt bold ALL CAPS (List of Figures, Abstract, etc.)."""
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    fmt(run, size=12, bold=True)

def sec_head(doc, text, level=1):
    style = {1:'Heading 2', 2:'Heading 3', 3:'Heading 4'}.get(level, 'Heading 3')
    size  = {1:14, 2:12, 3:12}.get(level, 12)
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    fmt(run, size=size, bold=True)

def fig_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    fmt(r, size=10, italic=True)

def add_img(doc, path, caption, w=13.0):
    if path and os.path.exists(path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.add_run().add_picture(path, width=Cm(w))
            fig_caption(doc, caption)
            return
        except Exception as e:
            print(f"  [WARN] Could not insert {os.path.basename(path)}: {e}")
    p = doc.add_paragraph(f'[PLACEHOLDER — {caption}]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.runs[0] if p.runs else p.add_run()
    fmt(r, size=11, italic=True, color='808080')
    fig_caption(doc, caption)

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(9)

def md_table(doc, rows_data):
    if not rows_data: return
    ncols = max(len(r) for r in rows_data)
    tbl = doc.add_table(rows=len(rows_data), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows_data):
        for ci in range(ncols):
            ct = row[ci].strip() if ci < len(row) else ''
            cell = tbl.cell(ri, ci)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            lo = ct.lower()
            is_hdr = (ri == 0)
            # determine cell color first so we can apply to runs
            if is_hdr:
                set_bg(cell, C['blue'][0]); txt_clr = C['blue'][1]
            elif lo == 'critical': set_bg(cell, C['critical'][0]); txt_clr = C['critical'][1]
            elif lo == 'high':     set_bg(cell, C['high'][0]);     txt_clr = C['high'][1]
            elif lo == 'medium':   set_bg(cell, C['medium'][0]);   txt_clr = None
            elif lo == 'low':      set_bg(cell, C['low'][0]);      txt_clr = C['low'][1]
            else:                                                   txt_clr = None
            add_inline_md(p, ct, size=11)
            for r in p.runs:
                r.font.size = Pt(11)
                if is_hdr or lo in ('critical','high','low'): r.font.bold = True
                if txt_clr: r.font.color.rgb = hex2rgb(txt_clr)
    doc.add_paragraph()

# ── Complete figure maps ──────────────────────────────────────────────────────

# ── Gantt / Action Plan tables ───────────────────────────────────────────────
#   on=1 → green cell (#92d050), off=0 → white
GANTT_T2 = {
    "caption": "Table 2: Action plan from February to March",
    "months":  [("February","4f9fd6","ffffff"), ("March","ffe600","000000")],
    "groups": [
        ("1","Environment Setup & Scope Definition",[
            ("Set up Kali Linux testing environment and tools",          [1,1,0,0,0,0,0,0]),
            ("Define engagement scope with Prestige Alliance Co., Ltd.", [1,1,0,0,0,0,0,0]),
            ("Obtain written authorization for neuralsh.com testing",    [0,1,1,0,0,0,0,0]),
        ]),
        ("2","Reconnaissance & Intelligence Gathering",[
            ("WHOIS analysis and DNS enumeration (A, MX, TXT, NS records)", [0,0,1,1,0,0,0,0]),
            ("Discover real origin IP via MX record and SSL certificate",    [0,0,1,1,0,0,0,0]),
            ("Subdomain enumeration and wildcard DNS detection",             [0,0,0,1,1,0,0,0]),
            ("JavaScript bundle analysis to extract API routes",             [0,0,0,0,1,1,0,0]),
        ]),
        ("3","Active Scanning & Enumeration",[
            ("Nmap port scan on origin server (103.16.62.217)",          [0,0,0,0,0,1,1,0]),
            ("Web directory enumeration with Dirb and Nikto",            [0,0,0,0,0,0,1,1]),
            ("Service fingerprinting — MySQL, SMTP, MikroTik, WHM",     [0,0,0,0,0,0,1,1]),
        ]),
    ]
}
GANTT_T3 = {
    "caption": "Table 3: Action plan from April to May",
    "months":  [("April","ffb84d","000000"), ("May","ffb3c6","000000")],
    "groups": [
        ("4","Vulnerability Assessment & Manual Testing",[
            ("Test authentication and rate limiting controls",              [1,1,0,0,0,0,0,0]),
            ("Analyze JWT token issuance and session management",           [1,1,0,0,0,0,0,0]),
            ("Assess exposed admin panels (WHM, cPanel, MikroTik, MySQL)", [0,1,1,0,0,0,0,0]),
            ("Assign CVSS v3.1 scores to all identified findings",          [0,0,1,1,0,0,0,0]),
        ]),
        ("5","Exploitation & Post-Exploitation",[
            ("Rate limit bypass via X-Forwarded-For header spoofing",      [0,0,0,1,1,0,0,0]),
            ("JWT token farming — 50 tokens, 0 rate-limit responses",      [0,0,0,1,1,0,0,0]),
            ("Attack chain mapping and lateral movement risk analysis",     [0,0,0,0,1,1,0,0]),
        ]),
        ("6","Report Writing & Thesis Documentation",[
            ("Write professional VAPT report with all findings and evidence",[0,0,0,0,0,1,1,0]),
            ("Design attack chain diagrams and infrastructure map",          [0,0,0,0,0,1,1,0]),
            ("Write thesis document and prepare for submission",             [0,0,0,0,0,0,1,1]),
        ]),
    ]
}

def gantt_table(doc, data):
    """Gantt table with vertical week headers, matching reference image style."""
    MONTH_COLORS = {"feb":"4472C4","mar":"FFC000","apr":"ED7D31","may":"FF99CC"}
    MONTH_TXT    = {"feb":"ffffff","mar":"000000","apr":"ffffff","may":"000000"}
    ACT_ON  = "70AD47"   # green for active weeks
    ACT_OFF = "ffffff"
    SEC_BG  = "D9D9D9"
    HDR_BG  = "1F3864"   # dark navy for No/Activities header

    months = data["months"]
    groups = data["groups"]
    total_rows = 2 + sum(1 + len(acts) for _,_,acts in groups)
    tbl = doc.add_table(rows=total_rows, cols=10)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Widths: No=0.7cm, Activities=6.0cm, 8 week cols=1.06cm each (≈16cm total)
    col_w = [Cm(0.7), Cm(6.0)] + [Cm(1.06)]*8
    for row in tbl.rows:
        for ci, w in enumerate(col_w):
            row.cells[ci].width = w

    def _write(cell, text, bg=None, txt_color="000000", bold=False,
                align=WD_ALIGN_PARAGRAPH.CENTER, size=9, vertical=False):
        """Clear cell, set background, write text with optional vertical rotation."""
        if bg:
            set_bg(cell, bg)
        set_cell_vert(cell, 'center')
        if vertical:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            td = OxmlElement('w:textDirection'); td.set(qn('w:val'), 'btLr')
            tcPr.append(td)
        for p in cell.paragraphs:   # clear any existing content
            for r in p.runs: r.clear()
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        if text:
            r = p.add_run(text)
            r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.bold = bold
            if txt_color != "000000": r.font.color.rgb = hex2rgb(txt_color)

    def _set_row_height(row, cm):
        tr = row._tr; trPr = tr.get_or_add_trPr()
        for x in trPr.findall(qn('w:trHeight')): trPr.remove(x)
        h = OxmlElement('w:trHeight')
        h.set(qn('w:val'), str(int(cm / 2.54 * 1440)))
        h.set(qn('w:hRule'), 'atLeast')
        trPr.append(h)

    m1, tc1, _ = months[0]; m2, tc2, _ = months[1]
    m1k = m1.lower()[:3]; m2k = m2.lower()[:3]
    m1bg = MONTH_COLORS.get(m1k, "4472C4"); m1tc = MONTH_TXT.get(m1k, "ffffff")
    m2bg = MONTH_COLORS.get(m2k, "FFC000"); m2tc = MONTH_TXT.get(m2k, "000000")

    # ── Row 0: No | Activities | Month1 (cols 2-5 merged) | Month2 (cols 6-9 merged)
    _write(tbl.cell(0,0), "No",         bg=HDR_BG, txt_color="ffffff", bold=True, size=10)
    _write(tbl.cell(0,1), "Activities", bg=HDR_BG, txt_color="ffffff", bold=True,
           align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
    tbl.cell(0,2).merge(tbl.cell(0,5))
    _write(tbl.cell(0,2), m1, bg=m1bg, txt_color=m1tc, bold=True, size=11)
    tbl.cell(0,6).merge(tbl.cell(0,9))
    _write(tbl.cell(0,6), m2, bg=m2bg, txt_color=m2tc, bold=True, size=11)

    # ── Row 1: blank | blank | Week 1–4 (vertical) | Week 1–4 (vertical)
    _write(tbl.cell(1,0), "", bg=HDR_BG)
    _write(tbl.cell(1,1), "", bg=HDR_BG)
    for wi in range(4):
        _write(tbl.cell(1, 2+wi), f"Week {wi+1}", bg=m1bg, txt_color=m1tc,
               bold=True, size=8, vertical=True)
        _write(tbl.cell(1, 6+wi), f"Week {wi+1}", bg=m2bg, txt_color=m2tc,
               bold=True, size=8, vertical=True)
    _set_row_height(tbl.rows[1], 2.2)   # tall enough for rotated "Week N" text

    # ── Data rows
    ri = 2
    for grp_no, grp_title, activities in groups:
        # Section header: number | title spanning remaining 9 cols
        _write(tbl.cell(ri,0), grp_no, bg=SEC_BG, bold=True, size=10)
        tbl.cell(ri,1).merge(tbl.cell(ri,9))
        _write(tbl.cell(ri,1), grp_title, bg=SEC_BG, bold=True,
               align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
        ri += 1
        for act_text, weeks in activities:
            _write(tbl.cell(ri,0), "", bg="F2F2F2")
            _write(tbl.cell(ri,1), act_text, align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
            for wi, on in enumerate(weeks):
                _write(tbl.cell(ri, 2+wi), "", bg=ACT_ON if on else ACT_OFF)
            ri += 1

    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_after = Pt(12)
    r3 = cap.add_run(data["caption"])
    r3.font.name = 'Times New Roman'; r3.font.size = Pt(10); r3.font.italic = True

# Logos / diagrams for Chapters 1–3 (key = figure number as string)
LOGO_FIGURES = {
    "1":  (os.path.join(IMGS, "prestige_logo.png"),        "Figure 1: Prestige Alliance company logo"),
    "2":  (os.path.join(IMGS, "map.png"),                   "Figure 2: Prestige Alliance office location"),
    "3":  (os.path.join(IMGS, "htb_logo_fixed.png"),         "Figure 3: Hack The Box platform logo"),
    "4":  (os.path.join(IMGS, "cpts_logo.png"),             "Figure 4: HTB Certified Penetration Testing Specialist (CPTS) logo"),
    "5":  (os.path.join(IMGS, "cwes_logo.png"),             "Figure 5: HTB Certified Web Exploitation Specialist (CWES) logo"),
    "6":  (os.path.join(IMGS, "pentest_flow.drawio.png"),   "Figure 6: Penetration testing process diagram"),
    "7":  (os.path.join(IMGS, "owasp_logo.png"),            "Figure 7: OWASP (Open Web Application Security Project) logo"),
    "8":  (os.path.join(IMGS, "nist_logo.png"),             "Figure 8: NIST SP 800-115 logo"),
    "9":  (os.path.join(IMGS, "cvss_logo.png"),             "Figure 9: CVSS v3.1 scoring system logo"),
    "10": (os.path.join(IMGS, "nmap_logo.png"),             "Figure 10: Nmap network scanner logo"),
    "11": (os.path.join(IMGS, "nikto_logo_fixed.png"),       "Figure 11: Nikto web server scanner logo"),
    "12": (os.path.join(IMGS, "shodan_logo_fixed.png"),      "Figure 12: Shodan internet-connected device search engine logo"),
    "13": (os.path.join(IMGS, "burpsuite_logo.png"),        "Figure 13: Burp Suite Community Edition logo"),
    "14": (os.path.join(IMGS, "cyberchef_logo.png"),        "Figure 14: CyberChef data analysis tool logo"),
    "15": (os.path.join(IMGS, "claude_logo.png"),           "Figure 15: Claude (Anthropic) AI assistant logo"),
}

# Chapter 4 screenshots (key = "4.N")
SHOT_FIGURES = {
    "4.1":  (os.path.join(SHOTS, "fig41_virtualbox.png"),                "Figure 4.1: VirtualBox showing Kali Linux VM running"),
    "4.2":  (os.path.join(SHOTS, "fig42_ssh_kali.png"),                  "Figure 4.2: SSH session from Linux Mint terminal connected to Kali VM"),
    "4.3":  (os.path.join(SHOTS, "fig43_tools_verify.png"),              "Figure 4.3: Tool version verification — Nmap, Nikto, Metasploit"),
    "4.4":  (os.path.join(SHOTS, "fig44_website.png"),                   "Figure 4.4: neuralsh.com homepage in browser"),
    "4.5":  (os.path.join(SHOTS, "fig45_whois.png"),                     "Figure 4.5: WHOIS lookup for neuralsh.com"),
    "4.6":  (os.path.join(SHOTS, "fig46_dig_mx.png"),                    "Figure 4.6: dig MX — origin IP discovery and DNS records"),
    "4.7":  (os.path.join(SHOTS, "fig47_wildcard_dns.png"),              "Figure 4.7: dig confirming wildcard DNS resolves to 103.16.62.217"),
    "4.8":  (os.path.join(SHOTS, "fig48_ssl_cert.png"),                  "Figure 4.8: SSL certificate mismatch — endoncambodia.com / onesala.com on origin"),
    "4.9":  (os.path.join(SHOTS, "fig49_nmap.png"),                      "Figure 4.9: Nmap full port scan of 103.16.62.217 — twenty-three open ports"),
    "4.10": (os.path.join(SHOTS, "fig410_cpanel_whm.png"),               "Figure 4.10: cPanel and WHM login pages accessible from the public internet"),
    "4.11": (os.path.join(SHOTS, "fig411_mikrotik.png"),                 "Figure 4.11: MikroTik RouterOS WebFig login page — publicly accessible"),
    "4.12": (os.path.join(SHOTS, "fig412_nikto.png"),                    "Figure 4.12: Nikto web scanner results against neuralsh.com"),
    "4.13": (os.path.join(SHOTS, "fig413_nuclei.png"),                   "Figure 4.13: Nuclei vulnerability scan — zero matches returned"),
    "4.14": (os.path.join(SHOTS, "fig414a_burp_cloudflare.png"),          "Figure 4.14: Burp Suite — Cloudflare WAF bypass confirmation (cloudflare vs Apache headers)"),
    "4.15": (os.path.join(SHOTS, "fig415b_ratelimit_bypass.png"),        "Figure 4.15: Rate limit bypass — fifty/fifty tokens, zero rate-limit responses"),
    "4.16": (os.path.join(SHOTS, "fig416_jwt_404.png"),                  "Figure 4.16: JWT endpoint confirmed live — HTTP 200 as of eleven June 2026"),
    "4.17": (os.path.join(SHOTS, "fig417_whm_browser.png"),              "Figure 4.17: WHM admin panel at 103.16.62.217:2087 — publicly accessible"),
    "4.18": (os.path.join(SHOTS, "fig418_mikrotik_browser.png"),         "Figure 4.18: MikroTik WebFig — RouterOS v6.49.18, admin pre-filled"),
    "4.19": (os.path.join(SHOTS, "fig419_mysql_connect.png"),            "Figure 4.19: MySQL connection from public IP — authentication prompt exposed"),
    "4.20": (os.path.join(SHOTS, "fig415_ratelimit_30requests.png"),     "Figure 4.20: Terminal showing thirty consecutive HTTP 200 responses — rate limiting not triggered"),
    "4.21": (os.path.join(SHOTS, "fig415_jwt_token_burp.png"),           "Figure 4.21: Burp Suite showing live JWT token response from /web/v1/init/token"),
    "4.22": (os.path.join(SHOTS, "fig422_attack_chains.png"),            "Figure 4.22: Attack chain diagram — four confirmed attack vectors"),
    # Chapter 5 figures
    "5.1":  (os.path.join(IMGS,  "recommended_architecture.png"),        "Figure 5.1: Recommended network security architecture"),
}

# Appendix screenshots
APPENDIX_SHOTS = {
    "A": [
        (os.path.join(SHOTS, "fig49_nmap.png"),
         "Figure A.1: Nmap full port scan — origin server (twenty-three open ports confirmed)"),
    ],
    "B": [
        (os.path.join(SHOTS, "fig48b_ssl_onesala.png"),
         "Figure B.1: SSL certificate showing onesala.com on same origin server — shared hosting confirmed"),
    ],
    "C": [
        (os.path.join(SHOTS, "fig415_ratelimit_30requests.png"),
         "Figure C.1: Thirty consecutive HTTP 200 responses — rate limiting not active"),
        (os.path.join(SHOTS, "fig415b_ratelimit_bypass.png"),
         "Figure C.2: Rate limit bypass confirmation — fifty tokens collected with zero rate-limit responses"),
    ],
    "D": [
        (os.path.join(SHOTS, "fig415_jwt_token_burp.png"),
         "Figure D.1: JWT token response captured in Burp Suite"),
        (os.path.join(SHOTS, "fig415b_jwt_token.png"),
         "Figure D.2: JWT token decoded — header and payload structure"),
        (os.path.join(SHOTS, "fig416_jwt_404.png"),
         "Figure D.3: JWT endpoint returning HTTP 200 — confirmed live and unpatched"),
    ],
    "E": [
        (os.path.join(SHOTS, "fig417_whm_browser.png"),
         "Figure E.1: WHM root admin panel accessible from public internet — no authentication bypass required"),
        (os.path.join(SHOTS, "fig418_mikrotik_browser.png"),
         "Figure E.2: MikroTik RouterOS admin panel accessible from public internet"),
        (os.path.join(SHOTS, "fig419_mysql_connect.png"),
         "Figure E.3: MySQL port 3306 accepting remote connections from external host"),
    ],
}

# ── Finding Card (SE_LYTHENG style) ──────────────────────────────────────────
def finding_card(doc, f):
    sev = f['severity'].lower()
    fill, txt = C[sev]
    sec_head(doc, f"{f['id']} — {f['title']}", level=2)
    tbl = doc.add_table(rows=8, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in tbl.rows:
        row.cells[0].width = Cm(3.4)
        row.cells[1].width = Cm(5.8)
        row.cells[2].width = Cm(5.8)
    # Row 0 — badge
    c0 = tbl.cell(0,0); c0.merge(tbl.cell(0,2))
    badge = f"{f['severity'].upper()}   CVSS {f['cvss']}   {f['id']}: {f['title']}"
    if f.get('status') == 'Exploited': badge += "   [CONFIRMED EXPLOITED]"
    p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.first_line_indent = Cm(0)
    r0 = p0.add_run(badge); fmt(r0, size=11, bold=True, color=txt)
    set_bg(c0, fill); set_cell_vert(c0)
    def lv(ri, lbl, val):
        lc = tbl.cell(ri,0); vc = tbl.cell(ri,1); vc.merge(tbl.cell(ri,2))
        set_bg(lc, C['label'][0]); set_cell_vert(lc)
        lc.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
        vc.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
        fmt(lc.paragraphs[0].add_run(lbl), size=11, bold=True)
        fmt(vc.paragraphs[0].add_run(val), size=11)
    def ll(ri, lbl, items):
        lc = tbl.cell(ri,0); vc = tbl.cell(ri,1); vc.merge(tbl.cell(ri,2))
        set_bg(lc, C['label'][0]); set_cell_vert(lc)
        lc.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
        fmt(lc.paragraphs[0].add_run(lbl), size=11, bold=True)
        for i, item in enumerate(items):
            p_ = vc.paragraphs[0] if i==0 else vc.add_paragraph()
            p_.paragraph_format.first_line_indent = Cm(0)
            fmt(p_.add_run(f"{i+1}.  {item}"), size=11)
    lv(1, "Description:", f['description'])
    lc2=tbl.cell(2,0); lk=tbl.cell(2,1); rk=tbl.cell(2,2)
    set_bg(lc2, C['label'][0]); set_cell_vert(lc2)
    lc2.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
    fmt(lc2.paragraphs[0].add_run("Risks:"), size=11, bold=True)
    lsev=f['likelihood'].lower(); lf,lt=C.get(lsev,C['medium'])
    set_bg(lk,lf); set_cell_vert(lk)
    plk=lk.paragraphs[0]; plk.alignment=WD_ALIGN_PARAGRAPH.CENTER
    plk.paragraph_format.first_line_indent=Cm(0)
    fmt(plk.add_run(f"Likelihood: {f['likelihood']}"), size=11, bold=True, color=lt)
    set_bg(rk,fill); set_cell_vert(rk)
    prk=rk.paragraphs[0]; prk.alignment=WD_ALIGN_PARAGRAPH.CENTER
    prk.paragraph_format.first_line_indent=Cm(0)
    fmt(prk.add_run(f"Risk Rating: {f['risk_rating']}"), size=11, bold=True, color=txt)
    lv(3,"Impact:",      f['impact'])
    lv(4,"Tool Used:",   f['tool'])
    ll(5,"References:",  f['references'])
    lv(6,"Remediation:", f['remediation'])
    lv(7,"Evidence:",    f['evidence'])
    doc.add_paragraph()

# ── All 20 findings ───────────────────────────────────────────────────────────
FINDINGS = [
  {"id":"N-001","severity":"Critical","cvss":"9.8","risk_rating":"Critical",
   "title":"MySQL Port 3306 Exposed to Internet",
   "host":"103.16.62.217:3306","tool":"Nmap, mysql-client","likelihood":"High","status":"Confirmed",
   "description":"The MySQL 8.0.43 database server is directly accessible from the public internet on port 3306, accepting full TCP connections from any external IP address. Authentication can be attempted directly at the database layer, bypassing all application-level controls.",
   "impact":"Successful access grants full read and write access to all application databases including user records, session data, and API secrets. MySQL UDF features could be abused for remote command execution.",
   "references":["https://owasp.org/Top10/A05_2021-Security_Misconfiguration/","https://cve.mitre.org/cgi-bin/cvename.cgi?name=CVE-2012-2122"],
   "remediation":"Restrict port 3306 to localhost or a trusted management VPN via firewall rules. Database ports must never be exposed to the public internet.",
   "evidence":"Figure 4.9 — Nmap scan confirming 3306/tcp open mysql MySQL 8.0.43; Figure 4.19 — MySQL connection accepted from public IP"},
  {"id":"N-002","severity":"Critical","cvss":"9.1","risk_rating":"Critical",
   "title":"MikroTik Admin Panel Exposed (Port 9001)",
   "host":"103.16.62.217:9001","tool":"Nmap, curl, browser","likelihood":"High","status":"Confirmed",
   "description":"The MikroTik RouterOS web configuration interface (WebFig) is publicly accessible on port 9001, returning HTTP 200 with a fully rendered login page. The login form pre-fills admin as the default username. MikroTik devices ship with a factory default of admin/blank password.",
   "impact":"An attacker with RouterOS access can modify routing tables, capture all network traffic, create persistent backdoor accounts, remove firewall rules, and disrupt all hosted services.",
   "references":["https://owasp.org/Top10/A05_2021-Security_Misconfiguration/","https://nvd.nist.gov/vuln/detail/CVE-2018-14847"],
   "remediation":"Restrict port 9001 to a trusted management IP via firewall. Change MikroTik default credentials immediately. Disable unused RouterOS services.",
   "evidence":"Figure 4.11 — curl response; Figure 4.18 — Browser showing MikroTik WebFig login, RouterOS v6.49.18, admin pre-filled"},
  {"id":"N-013","severity":"Critical","cvss":"9.8","risk_rating":"Critical",
   "title":"cPanel Admin Panel Exposed (Port 2083)",
   "host":"103.16.62.217:2083","tool":"Nmap, curl","likelihood":"High","status":"Confirmed",
   "description":"The cPanel web hosting control panel is publicly accessible on port 2083 with no IP restriction and no two-factor authentication. cPanel provides full management of the hosting account including file manager, phpMyAdmin, email control, FTP credentials, and subdomain configuration.",
   "impact":"An attacker who compromises the registered email account can reset the cPanel password via the login page reset link and gain full hosting account access without brute force.",
   "references":["https://owasp.org/Top10/A05_2021-Security_Misconfiguration/","https://docs.cpanel.net/cpanel/security/two-factor-authentication-for-cpanel/"],
   "remediation":"Restrict port 2083 to trusted management IPs via Cloudways firewall. Enable two-factor authentication on cPanel immediately.",
   "evidence":"Figure 4.10 — curl response showing cPanel Login title confirming public accessibility"},
  {"id":"N-014","severity":"Critical","cvss":"10.0","risk_rating":"Critical",
   "title":"WHM Root Panel Exposed (Port 2087)",
   "host":"103.16.62.217:2087","tool":"Nmap, curl, browser","likelihood":"High","status":"Confirmed",
   "description":"WHM (Web Host Manager) is the root-level server administration panel for cPanel-based servers. It is publicly accessible on port 2087 with no IP restriction, no two-factor authentication, and no account lockout. WHM provides a built-in Terminal, full server configuration, and access to every hosted account.",
   "impact":"An attacker who gains WHM access compromises not only neuralsh.com but every other hosted domain on the shared server simultaneously. CVSS score: 10.0 (maximum).",
   "references":["https://owasp.org/Top10/A05_2021-Security_Misconfiguration/","https://docs.cpanel.net/whm/security-center/"],
   "remediation":"Immediately firewall port 2087 to trusted IPs only. This single action eliminates the highest-risk vector in the engagement.",
   "evidence":"Figure 4.17 — Browser showing WHM login page at 103.16.62.217:2087 (SSL warning visible, no authentication bypass required)"},
  {"id":"N-003","severity":"High","cvss":"7.2","risk_rating":"High",
   "title":"SSH Port 22 Publicly Accessible",
   "host":"103.16.62.217:22","tool":"Nmap, ssh","likelihood":"Medium","status":"Confirmed",
   "description":"OpenSSH 8.9p1 is publicly accessible on port 22. SSH key-based authentication was confirmed as the enforced method for tested usernames. However, the port remains exposed to brute-force attempts for any account that may have password authentication enabled.",
   "impact":"The SMTP banner confirms this is a cPanel shared hosting server where the primary user holds sudo privileges. Successful SSH compromise is equivalent to root access to the entire server.",
   "references":["https://owasp.org/Top10/A05_2021-Security_Misconfiguration/","https://www.cisecurity.org/benchmark/ubuntu_linux"],
   "remediation":"Enforce PasswordAuthentication no in /etc/ssh/sshd_config. Install fail2ban. Restrict SSH to a trusted management IP via firewall.",
   "evidence":"Figure 4.9 — Nmap scan confirming 22/tcp open OpenSSH 8.9p1"},
  {"id":"N-004","severity":"High","cvss":"7.5","risk_rating":"High",
   "title":"Shared Hosting Lateral Movement Risk",
   "host":"103.16.62.217 (shared infrastructure)","tool":"openssl, curl","likelihood":"Medium","status":"Confirmed",
   "description":"The origin server is a shared cPanel hosting environment (Cloudways/cprapid.com) hosting multiple clients on the same physical server. SSL certificates issued to *.onesala.com and www.endoncambodia.com confirmed the presence of other tenants.",
   "impact":"A WHM-level compromise grants access to the files, databases, and email of every co-hosted account simultaneously, constituting a breach affecting organizations unrelated to neuralsh.com.",
   "references":["https://owasp.org/Top10/A05_2021-Security_Misconfiguration/"],
   "remediation":"Migrate neuralsh.com to a dedicated server or isolated cloud instance. Short-term: firewall WHM immediately.",
   "evidence":"Figure 4.8 — SSL certificate showing onesala.com and endoncambodia.com on the same origin server"},
  {"id":"N-005","severity":"High","cvss":"9.1","risk_rating":"High",
   "title":"Rate Limit Bypass via X-Forwarded-For Header Spoofing",
   "host":"neuralsh.com/web/v1/init/token","tool":"Burp Suite, curl, Python","likelihood":"High","status":"Exploited",
   "description":"The application's rate-limiting mechanism uses the client-supplied X-Forwarded-For header as the IP identifier. Because X-Forwarded-For is fully client-controlled, an attacker can supply a different random IP on each request, making every request appear to originate from a unique source.",
   "impact":"Fifty consecutive requests with rotating X-Forwarded-For values all received HTTP 200 — zero rate-limit responses. This bypass enables unlimited JWT token farming. STATUS: CONFIRMED EXPLOITED.",
   "references":["https://owasp.org/Top10/A04_2021-Insecure_Design/","https://developers.cloudflare.com/fundamentals/reference/http-headers/"],
   "remediation":"Replace X-Forwarded-For with Cloudflare's CF-Connecting-IP header for all rate-limiting logic. CF-Connecting-IP cannot be spoofed by the client.",
   "evidence":"Figure 4.15 — fifty/fifty tokens, zero rate-limit responses; Figure 4.20 — thirty consecutive HTTP 200 with no rate limiting"},
  {"id":"N-015","severity":"High","cvss":"7.5","risk_rating":"High",
   "title":"Webmail Interface Exposed (Port 2096)",
   "host":"103.16.62.217:2096","tool":"Nmap, curl","likelihood":"Medium","status":"Confirmed",
   "description":"The cPanel Webmail login interface is publicly accessible on port 2096 without IP restriction, allowing brute-force attacks against all @neuralsh.com email accounts.",
   "impact":"A compromised email account can trigger password resets for cPanel and application accounts, and can send phishing emails from a legitimate @neuralsh.com address with valid DKIM signatures.",
   "references":["https://owasp.org/Top10/A05_2021-Security_Misconfiguration/"],
   "remediation":"Restrict port 2096 to trusted IPs via firewall. Enforce strong passwords and account lockout on webmail.",
   "evidence":"Nmap scan — 2096/tcp open cPanel Webmail; Figure 4.9"},
  {"id":"N-018","severity":"High","cvss":"7.5","risk_rating":"High",
   "title":"Additional cPanel Interfaces Exposed (Ports 2078, 2091)",
   "host":"103.16.62.217:2078, 2091","tool":"Nmap, curl (follow-up, ten June 2026)","likelihood":"Medium","status":"Confirmed",
   "description":"Ports 2078 and 2091 were not present in the June six baseline but appeared open in the June ten follow-up scan, both returning HTTP 401 Basic Auth. Their appearance indicates infrastructure changes without security review.",
   "impact":"Additional credential brute-force attack surface for server management interfaces. Absence of change management controls confirmed.",
   "references":["https://owasp.org/Top10/A05_2021-Security_Misconfiguration/"],
   "remediation":"Restrict ports 2078 and 2091 via firewall. Conduct a full port audit and close all ports without documented business purpose.",
   "evidence":"Follow-up Nmap scan (ten June 2026) — 2078/tcp and 2091/tcp open returning HTTP 401"},
  {"id":"N-006","severity":"Medium","cvss":"5.3","risk_rating":"Medium",
   "title":"API Routes Exposed in Client-Side JavaScript",
   "host":"neuralsh.com/_nuxt/","tool":"curl, grep","likelihood":"Low","status":"Confirmed",
   "description":"The complete API route structure is disclosed in the compiled Nuxt.js JavaScript bundle. Routes including /web/v1/init/token, /web/v1/text/search, /web/v1/image/search, /web/v1/report/save, and /api/geocode were extracted without authentication.",
   "impact":"This eliminated the need for directory brute force and directly enabled targeted exploitation of the unauthenticated token endpoint (N-007). All internal API endpoints are effectively public.",
   "references":["https://owasp.org/Top10/A01_2021-Broken_Access_Control/"],
   "remediation":"Minimize route exposure in the bundle. Implement an API gateway with strict allowlisting. Secure all endpoints regardless of client-side exposure.",
   "evidence":"curl extraction of API routes from /_nuxt/Bpuv52g-.js bundle"},
  {"id":"N-007","severity":"Medium","cvss":"5.3","risk_rating":"Medium",
   "title":"Unauthenticated JWT Token Issuance",
   "host":"neuralsh.com/web/v1/init/token","tool":"curl, CyberChef, hashcat, Burp Suite","likelihood":"High","status":"Exploited",
   "description":"The token endpoint issues a valid signed JWT token to any unauthenticated caller with no credentials, API key, or device fingerprint required. Tokens carry type:guest and a thirty-minute validity window. Confirmed still live as of eleven June 2026.",
   "impact":"Fifty valid tokens farmed in approximately three seconds using the N-005 bypass. Each token granted real API access. JWT secret not cracked (rockyou.txt exhausted). STATUS: CONFIRMED EXPLOITED.",
   "references":["https://owasp.org/Top10/A07_2021-Identification_and_Authentication_Failures/","https://portswigger.net/web-security/jwt"],
   "remediation":"Require Cloudflare Turnstile challenge or browser fingerprint before issuing tokens. Implement token binding.",
   "evidence":"Figure 4.16 — endpoint live HTTP 200; Figure 4.21 — Burp showing JWT response; Appendix D"},
  {"id":"N-008","severity":"Medium","cvss":"5.9","risk_rating":"Medium",
   "title":"SSL Certificate Mismatch on Backend Server",
   "host":"103.16.62.217:443, 110, 143, 465","tool":"openssl","likelihood":"Low","status":"Confirmed",
   "description":"The SSL certificate served on the backend HTTPS port is issued to www.endoncambodia.com, not neuralsh.com. Mail certificates are issued to *.onesala.com. Any direct connection to the origin IP triggers a certificate warning.",
   "impact":"Confirms infrastructure-level SSL hygiene has not been maintained and corroborates the shared hosting risk (N-004). Co-tenant domain names are disclosed via certificate inspection.",
   "references":["https://owasp.org/Top10/A02_2021-Cryptographic_Failures/"],
   "remediation":"Issue a dedicated SSL certificate for neuralsh.com at the origin level, or use a multi-SAN certificate covering all hosted domains.",
   "evidence":"Figure 4.8 — openssl output showing www.endoncambodia.com and *.onesala.com"},
  {"id":"N-009","severity":"Medium","cvss":"5.4","risk_rating":"Medium",
   "title":"CSP Allows unsafe-inline Scripts",
   "host":"neuralsh.com","tool":"curl","likelihood":"Medium","status":"Confirmed",
   "description":"The Content-Security-Policy header includes 'unsafe-inline' in script-src, allowing execution of inline JavaScript including dynamically injected script blocks and event handlers.",
   "impact":"If any XSS vulnerability is present, the CSP would not prevent exploitation of inline injection, rendering the entire CSP protection ineffective.",
   "references":["https://owasp.org/Top10/A03_2021-Injection/","https://developer.mozilla.org/en-US/docs/Web/HTTP/CSP"],
   "remediation":"Replace 'unsafe-inline' with per-request nonces. Nuxt.js has built-in CSP nonce support via its security plugin.",
   "evidence":"curl response headers: Content-Security-Policy: script-src 'self' 'unsafe-inline'"},
  {"id":"N-010","severity":"Medium","cvss":"4.3","risk_rating":"Medium",
   "title":"Wildcard DNS Configured",
   "host":"*.neuralsh.com","tool":"dig","likelihood":"Low","status":"Confirmed",
   "description":"A wildcard DNS record causes any subdomain of neuralsh.com to resolve to 103.16.62.217, enabling use of subdomains like login.neuralsh.com in phishing campaigns.",
   "impact":"Wildcard subdomains resolve to a real IP address, significantly increasing the credibility of phishing attacks using the neuralsh.com brand.",
   "references":["https://owasp.org/Top10/A05_2021-Security_Misconfiguration/"],
   "remediation":"Remove the wildcard DNS record from Cloudflare. Define only explicit A records for legitimate subdomains.",
   "evidence":"Figure 4.7 — dig showing randomxyz123.neuralsh.com resolves to 103.16.62.217"},
  {"id":"N-016","severity":"Medium","cvss":"5.3","risk_rating":"Medium",
   "title":"Directory Listing Enabled on Apache",
   "host":"103.16.62.217","tool":"curl","likelihood":"Low","status":"Confirmed",
   "description":"Apache directory listing is enabled at the document root. A direct HTTP request to the origin IP returns an 'Index of /' page listing all files with names and modification timestamps.",
   "impact":"Any deployed files would be visible to an attacker, potentially exposing backup files or configuration files that accelerate targeted attacks.",
   "references":["https://owasp.org/Top10/A05_2021-Security_Misconfiguration/"],
   "remediation":"Add 'Options -Indexes' to the Apache configuration or .htaccess file.",
   "evidence":"curl response showing 'Index of /' from 103.16.62.217"},
  {"id":"N-020","severity":"Medium","cvss":"6.1","risk_rating":"Medium",
   "title":"Shared Hosting Co-Tenant Identified: onesala.com",
   "host":"103.16.62.217","tool":"curl, openssl","likelihood":"Medium","status":"Confirmed",
   "description":"HTTP redirects from ports 2077 and 2082 explicitly redirect to www.onesala.com:2078 and www.onesala.com:2083, confirming onesala.com as a named, verifiable co-tenant on the same physical server.",
   "impact":"Elevates N-004 from theoretical to confirmed with a real third-party victim. A breach of neuralsh.com now carries documented third-party data protection implications.",
   "references":["https://owasp.org/Top10/A05_2021-Security_Misconfiguration/"],
   "remediation":"Long-term: migrate to isolated infrastructure. Short-term: firewall WHM immediately.",
   "evidence":"curl showing HTTP redirect from 103.16.62.217:2077 to www.onesala.com:2078"},
  {"id":"N-011","severity":"Low","cvss":"3.7","risk_rating":"Low",
   "title":"Information Disclosure via Error Messages",
   "host":"neuralsh.com/api/*","tool":"curl","likelihood":"Low","status":"Confirmed",
   "description":"API error responses disclose internal details including full request URLs, internal status messages, parameter names, and framework identification strings.",
   "impact":"Verbose error messages reduce reconnaissance effort required to map the API surface and understand expected parameter formats.",
   "references":["https://owasp.org/Top10/A05_2021-Security_Misconfiguration/"],
   "remediation":"Return generic error responses. Log details server-side only. Remove parameter names from client-facing errors.",
   "evidence":"curl to /api/geocode with missing parameters returning detailed JSON error"},
  {"id":"N-012","severity":"Low","cvss":"3.1","risk_rating":"Low",
   "title":"SPF Softfail / DMARC Quarantine",
   "host":"neuralsh.com DNS","tool":"dig","likelihood":"Low","status":"Confirmed",
   "description":"SPF uses ~all (softfail) rather than -all (hardfail). DMARC policy is p=quarantine rather than p=reject. Spoofed @neuralsh.com emails are flagged but not rejected.",
   "impact":"Phishing emails using the neuralsh.com brand may be delivered to recipients whose mail servers do not strictly enforce SPF.",
   "references":["https://dmarc.org/","https://tools.ietf.org/html/rfc7208"],
   "remediation":"Change SPF to: v=spf1 +mx +a +ip4:103.16.62.217 -all. Change DMARC to: v=DMARC1; p=reject; rua=mailto:dmarc@neuralsh.com.",
   "evidence":"Figure 4.6 — dig TXT showing SPF softfail and DMARC p=quarantine"},
  {"id":"N-017","severity":"Low","cvss":"3.1","risk_rating":"Low",
   "title":"cPanel Version Disclosure",
   "host":"103.16.62.217:2083","tool":"curl","likelihood":"Low","status":"Confirmed",
   "description":"The cPanel version is disclosed via magic revision numbers embedded in static asset paths: cPanel_magic_revision_1698766296. This timestamp can be cross-referenced to identify the exact cPanel version.",
   "impact":"Version disclosure enables targeted CVE lookup for that specific cPanel version without active version scanning.",
   "references":["https://owasp.org/Top10/A05_2021-Security_Misconfiguration/"],
   "remediation":"Update cPanel to the latest stable version. Configure cPanel to suppress magic revision numbers.",
   "evidence":"curl to 103.16.62.217:2083 — HTML source contains cPanel_magic_revision_1698766296"},
  {"id":"N-019","severity":"Low","cvss":"3.5","risk_rating":"Low",
   "title":"SMTP Port 25 Transitioned from Filtered to Open",
   "host":"103.16.62.217:25","tool":"Nmap, nc (follow-up, ten June 2026)","likelihood":"Low","status":"Confirmed",
   "description":"Port 25 was filtered in the June six baseline scan and open in the June ten follow-up, indicating a firewall rule was removed without a documented change control process.",
   "impact":"An open SMTP port enables mail relay testing, SMTP user enumeration, and potential spam relay abuse.",
   "references":["https://owasp.org/Top10/A05_2021-Security_Misconfiguration/"],
   "remediation":"Investigate why the port twenty-five firewall rule was removed. Restore the firewall block if direct SMTP delivery is not required.",
   "evidence":"Comparison: June six Nmap baseline (25/tcp filtered) vs June ten follow-up (25/tcp open)"},
]

# ── Front matter (LOF / LOT / LOA as plain text, not tables) ──────────────────
def lof_entry(doc, num, desc):
    """Single List of Figures entry — plain text, no table"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(f"Figure {num}:")
    fmt(r, size=11, bold=True)
    r2 = p.add_run(f"  {desc}")
    fmt(r2, size=11)

def lot_entry(doc, num, desc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(f"Table {num}:")
    fmt(r, size=11, bold=True)
    r2 = p.add_run(f"  {desc}")
    fmt(r2, size=11)

def loa_entry(doc, abbr, full):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(f"{abbr:<12}")
    fmt(r, size=11, bold=True)
    r2 = p.add_run(full)
    fmt(r2, size=11)

def add_front_matter(doc):
    # Title page
    def centre(text, size=12, bold=False, sa=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(sa)
        r = p.add_run(text); fmt(r, size=size, bold=bold)

    logo = os.path.join(IMGS, "itc_logo.png")
    if os.path.exists(logo):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0); p.paragraph_format.space_after = Pt(12)
        p.add_run().add_picture(logo, width=Cm(3))

    centre("INSTITUTION OF TECHNOLOGY OF CAMBODIA", 14, bold=True, sa=4)
    centre("Department of General Studies and Natural Sciences", 12, sa=4)
    centre("Bachelor of Information and Communication Engineering", 12, sa=24)
    centre("INTERNSHIP REPORT", 16, bold=True, sa=12)
    centre("VULNERABILITY ASSESSMENT AND PENETRATION TESTING (VAPT)", 14, bold=True, sa=6)
    centre("ON A WEB APPLICATION", 14, bold=True, sa=6)
    centre("A Case Study of neuralsh.com", 12, sa=24)

    tbl = doc.add_table(rows=5, cols=2); tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri,(lbl,val) in enumerate([
        ("Student Name:","[Your Full Name]"),("Student ID:","[Your Student ID]"),
        ("Academic Supervisor:","[Supervisor Name]"),
        ("Company Supervisor:","Prestige Alliance Co., Ltd."),("Submission Date:","June 2026"),
    ]):
        for ci,(txt,bld) in enumerate([(lbl,True),(val,False)]):
            c = tbl.cell(ri,ci); c.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
            fmt(c.paragraphs[0].add_run(txt), size=12, bold=bld)
    doc.add_paragraph()

    # ACKNOWLEDGMENT
    front_head(doc, "ACKNOWLEDGMENT")
    body(doc, "Before I start this report, I would like to express my sincere gratitude to our school and to everyone who has contributed to the making of this report. Everyone has dedicated their hard work and efforts, without which it would not have been possible to achieve this accomplishment. I am very grateful for the opportunity and the guidance that allowed me to grow and enhance my knowledge in both hard and soft skills. I would like to extend my appreciation to the following individuals:")
    body(doc, "First of all, I would like to offer my sincere thanks to **H.E. PO Kimtho, Director of the Institute of Technology of Cambodia, a.k.a (ITC)**. For his leadership and efforts that steer our institute to be one of the top universities in Cambodia. His cooperation with local and international enterprises has been instrumental in making this opportunity possible.")
    body(doc, "Secondly, I would like to express my appreciation to **MR. LAY Heng, Head of Department of Information and Communication Engineering (DICE) at ITC**. His guidance and management have provided excellent education and courses within our department which are the foundation of this internship.")
    body(doc, "Thirdly, I would like to give my deep gratitude to my supervisor, **DR. KUY Movsun, Lecturer at the Department of Information and Communication Engineering (DICE), ITC**. His kindness and dedication of time have guided this internship to achieve its goals. From the beginning of the internship to the final report, his thoughtful suggestions were invaluable.")
    body(doc, "Finally, I would like to express my sincere thanks to my advisor, **MR. KIM Sereyvuth, Head of Red Team at Prestige Alliance Co., Ltd**. He has allowed me to conduct this vulnerability assessment and penetration testing engagement and generously dedicated his invaluable time to allow me to learn more in cybersecurity skills. He provided the incredible opportunity that enhanced my technical skills.")
    body(doc, "Additionally, I would like to extend my deep appreciation to all others who have supported me during this internship. Your kindness and assistance have been essential to this accomplishment. I wish everyone all the best.")

    # KHMER ABSTRACT (មូលដ្ឋានទេច)
    p_kh = doc.add_paragraph(style='Heading 1')
    p_kh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rk = p_kh.add_run("មូលដ្ឋានទេច")
    rk.font.name = "Khmer OS"; rk.font.size = Pt(12)
    rk.font.bold = True; rk.font.color.rgb = RGBColor(0,0,0)
    p_kh_body = doc.add_paragraph()
    p_kh_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_kh_body.paragraph_format.first_line_indent = Cm(1.27)
    p_kh_body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p_kh_body.paragraph_format.space_after = Pt(6)
    r_kh = p_kh_body.add_run("[បញ្ចូលសរុបខ្លឹមសារជាភាសាខ្មែររបស់អ្នកនៅទីនេះ — paste your Khmer summary here]")
    r_kh.font.name = "Khmer OS"; r_kh.font.size = Pt(12)
    r_kh.font.italic = True; r_kh.font.color.rgb = RGBColor(0x80,0x80,0x80)

    # RÉSUMÉ (French)
    front_head(doc, "RÉSUMÉ")
    body(doc, "Ce rapport présente un engagement complet d'Évaluation des Vulnérabilités et de Tests de Pénétration (VAPT) réalisé sur neuralsh.com, une plateforme de recherche neurale alimentée par l'intelligence artificielle, dans le cadre d'un stage au sein de Prestige Alliance Co., Ltd. L'évaluation a été autorisée par écrit avant le début des tests, et toutes les activités ont été menées conformément au guide de test OWASP et aux normes éthiques établies.")
    body(doc, "Une méthodologie boîte noire a été appliquée, simulant un attaquant externe sans connaissance préalable du système cible. L'engagement a suivi le cycle complet des tests d'intrusion : reconnaissance, analyse, énumération, exploitation, post-exploitation et rapport. Vingt vulnérabilités distinctes ont été découvertes : quatre Critiques, cinq Élevées, sept Moyennes et quatre Faibles. Les résultats critiques comprennent une base de données MySQL exposée publiquement, un panneau d'administration de routeur réseau MikroTik, un panneau de contrôle d'hébergement cPanel et un panneau d'administration de serveur racine WHM, tous accessibles depuis l'internet public sans restriction d'adresse IP.")
    body(doc, "Un contrôle de limitation de débit a été contourné avec succès via la falsification d'en-têtes HTTP, permettant une génération illimitée de jetons API. Une analyse de vérification de suivi a confirmé que toutes les découvertes originales n'étaient pas corrigées et a identifié trois nouvelles vulnérabilités.")
    body(doc, "Mots-clés : Tests de pénétration, VAPT, Sécurité des applications web, OWASP, Contournement de limitation de débit, Analyse JWT, Exposition cPanel, MikroTik, Tests boîte noire, CVSS v3.1")

    # ABSTRACT
    front_head(doc, "ABSTRACT")
    body(doc, "This report presents a comprehensive Vulnerability Assessment and Penetration Testing (VAPT) engagement conducted on neuralsh.com, an AI-powered neural search platform, as part of an internship at Prestige Alliance Co., Ltd. The assessment was authorized in writing prior to testing, and all activities were conducted in accordance with the OWASP Testing Guide and established ethical hacking standards.")
    body(doc, "A black-box methodology was applied, simulating an external attacker with no prior knowledge of the target system. The engagement followed the complete penetration testing lifecycle: reconnaissance, scanning, enumeration, exploitation, post-exploitation, and reporting. Twenty distinct vulnerabilities were discovered: four Critical, five High, seven Medium, and four Low severity findings. Critical findings include a publicly exposed MySQL database, a MikroTik network router administration panel, a cPanel hosting control panel, and a WHM root server administration panel — all accessible from the public internet without IP restriction.")
    body(doc, "A rate-limiting control was successfully bypassed via HTTP header spoofing, enabling unlimited API token generation. A follow-up verification scan conducted on ten June 2026 confirmed all original findings remained unpatched and identified three additional findings.")
    body(doc, "Keywords: Penetration Testing, VAPT, Web Application Security, OWASP, Rate Limiting Bypass, JWT Analysis, cPanel Exposure, MikroTik, Black Box Testing, CVSS v3.1")

    # TABLE OF CONTENTS
    front_head(doc, "TABLE OF CONTENTS")
    add_toc_field(doc)

    # LIST OF FIGURES — plain text, no table
    front_head(doc, "LIST OF FIGURES")
    lof_entries = [
        ("1",    "Prestige Alliance company logo"),
        ("2",    "Prestige Alliance office location map"),
        ("3",    "Hack The Box platform logo"),
        ("4",    "HTB Certified Penetration Testing Specialist (CPTS) logo"),
        ("5",    "HTB Certified Web Exploitation Specialist (CWES) logo"),
        ("6",    "Penetration testing process diagram"),
        ("7",    "OWASP logo"),
        ("8",    "NIST SP 800-115 logo"),
        ("9",    "CVSS v3.1 scoring system logo"),
        ("10",   "Nmap network scanner logo"),
        ("11",   "Nikto web server scanner logo"),
        ("12",   "Shodan search engine logo"),
        ("13",   "Burp Suite Community Edition logo"),
        ("14",   "CyberChef data analysis tool logo"),
        ("15",   "Claude (Anthropic) AI assistant logo"),
        ("4.1",  "VirtualBox showing Kali Linux VM running"),
        ("4.2",  "SSH session from Linux Mint terminal connected to Kali VM"),
        ("4.3",  "Tool version verification — Nmap, Nikto, Metasploit"),
        ("4.4",  "neuralsh.com homepage in browser"),
        ("4.5",  "WHOIS lookup for neuralsh.com"),
        ("4.6",  "dig MX — origin IP discovery and DNS records"),
        ("4.7",  "Wildcard DNS — randomxyz123.neuralsh.com resolves to 103.16.62.217"),
        ("4.8",  "SSL certificate mismatch — endoncambodia.com and onesala.com on origin server"),
        ("4.9",  "Nmap full port scan of 103.16.62.217 — twenty-three open ports"),
        ("4.10", "cPanel and WHM login pages accessible from the public internet"),
        ("4.11", "MikroTik RouterOS WebFig login page — publicly accessible"),
        ("4.12", "Nikto web scanner results against neuralsh.com"),
        ("4.13", "Nuclei vulnerability scan — zero matches returned"),
        ("4.14", "Burp Suite — Cloudflare WAF bypass confirmation"),
        ("4.15", "Rate limit bypass — fifty/fifty tokens, zero rate-limit responses"),
        ("4.16", "JWT endpoint confirmed live — HTTP 200 as of eleven June 2026"),
        ("4.17", "WHM admin panel at 103.16.62.217:2087 — publicly accessible"),
        ("4.18", "MikroTik WebFig — RouterOS v6.49.18, admin pre-filled"),
        ("4.19", "MySQL connection from public IP — authentication prompt exposed"),
        ("4.20", "Thirty consecutive HTTP 200 responses — rate limiting not triggered"),
        ("4.21", "Burp Suite showing live JWT token response from /web/v1/init/token"),
        ("4.22", "Attack chain diagram — four confirmed attack vectors"),
        ("A.1",  "Nmap full port scan output — Appendix A"),
        ("C.1",  "Thirty requests, zero rate-limit responses — Appendix C"),
        ("C.2",  "Rate limit bypass confirmation — Appendix C"),
        ("D.1",  "JWT token response in Burp Suite — Appendix D"),
        ("D.2",  "JWT endpoint returning HTTP 200, eleven June 2026 — Appendix D"),
    ]
    for num, desc in lof_entries:
        lof_entry(doc, num, desc)

    # LIST OF TABLES — plain text
    front_head(doc, "LIST OF TABLES")
    lot_entries = [
        ("4.1",  "Testing tools used in this engagement"),
        ("5.1",  "Complete findings register — twenty vulnerabilities"),
        ("5.2",  "Severity distribution and CVSS score summary"),
        ("5.3",  "Remediation priority roadmap"),
    ]
    for num, desc in lot_entries:
        lot_entry(doc, num, desc)

    # LIST OF ABBREVIATIONS — plain text sorted A–Z
    front_head(doc, "LIST OF ABBREVIATIONS")
    abbrevs = [
        ("API",   "Application Programming Interface"),
        ("CDN",   "Content Delivery Network"),
        ("CORS",  "Cross-Origin Resource Sharing"),
        ("CSP",   "Content Security Policy"),
        ("CVSS",  "Common Vulnerability Scoring System"),
        ("CWE",   "Common Weakness Enumeration"),
        ("CWES",  "Certified Web Exploitation Specialist (HTB)"),
        ("DMARC", "Domain-based Message Authentication, Reporting and Conformance"),
        ("DNS",   "Domain Name System"),
        ("HTTP",  "Hypertext Transfer Protocol"),
        ("HTTPS", "Hypertext Transfer Protocol Secure"),
        ("ITC",   "Institution of Technology of Cambodia"),
        ("JWT",   "JSON Web Token"),
        ("NIST",  "National Institute of Standards and Technology"),
        ("OSINT", "Open Source Intelligence"),
        ("OWASP", "Open Web Application Security Project"),
        ("SMTP",  "Simple Mail Transfer Protocol"),
        ("SPF",   "Sender Policy Framework"),
        ("SQL",   "Structured Query Language"),
        ("SSL",   "Secure Sockets Layer"),
        ("TCP",   "Transmission Control Protocol"),
        ("TLS",   "Transport Layer Security"),
        ("URL",   "Uniform Resource Locator"),
        ("VAPT",  "Vulnerability Assessment and Penetration Testing"),
        ("WAF",   "Web Application Firewall"),
        ("WHM",   "Web Host Manager"),
        ("XFF",   "X-Forwarded-For (HTTP header)"),
        ("XSS",   "Cross-Site Scripting"),
    ]
    for abbr, full in abbrevs:
        loa_entry(doc, abbr, full)

    add_section_break(doc, restart_arabic=True)

# ── Markdown → docx ───────────────────────────────────────────────────────────
def process_md(doc):
    with open(MD, encoding='utf-8') as fh:
        lines = fh.readlines()

    in_code=False; code_buf=[]; in_table=False; table_rows=[]; in_53=False
    in_appendix=False; current_appendix=None

    i = 0
    while i < len(lines):
        raw = lines[i]; line = raw.rstrip('\n'); stripped = line.strip()

        # code blocks
        if stripped.startswith('```'):
            if not in_code: in_code=True; code_buf=[]
            else:
                in_code=False
                if not in_53: code_block(doc, '\n'.join(code_buf))
            i+=1; continue
        if in_code: code_buf.append(line); i+=1; continue

        if stripped in ('---','***','___'): i+=1; continue

        # markdown tables
        if stripped.startswith('|') and stripped.endswith('|'):
            cols = [c.strip() for c in stripped[1:-1].split('|')]
            if not all(re.fullmatch(r':?-+:?',c.replace(' ','')) for c in cols):
                table_rows.append(cols)
            in_table=True; i+=1; continue
        else:
            if in_table:
                if not in_53: md_table(doc, table_rows)
                table_rows=[]; in_table=False

        # ── Action plan Gantt tables ──────────────────────────────────────
        if re.match(r'\*\[Table 2:.*ACTION_PLAN', stripped):
            gantt_table(doc, GANTT_T2); i+=1; continue
        if re.match(r'\*\[Table 3:.*ACTION_PLAN', stripped):
            gantt_table(doc, GANTT_T3); i+=1; continue

        # ── Figure insertions ──────────────────────────────────────────────
        # Logo figures: *[Figure N: description — filename.ext]*
        m_logo = re.match(r'\*\[Figure (\d+): ([^—\]]+) — ([^\]]+)\]\*', stripped)
        if m_logo:
            num = m_logo.group(1); desc = m_logo.group(2).strip()
            fname = m_logo.group(3).strip()
            if num in LOGO_FIGURES:
                path, cap = LOGO_FIGURES[num]
                add_img(doc, path, cap, w=8.0)
            else:
                path = os.path.join(IMGS, fname)
                add_img(doc, path, f"Figure {num}: {desc}", w=8.0)
            i+=1; continue

        # Screenshot figures: *[Figure N.N: description ... — screenshot]*
        m_shot = re.match(r'\*\[Figure (\d+\.\d+): (.+) — screenshot\]\*', stripped)
        if m_shot:
            num = m_shot.group(1)
            if num in SHOT_FIGURES:
                path, cap = SHOT_FIGURES[num]
                add_img(doc, path, cap, w=14.0)
            i+=1; continue

        # Image figures: *[Figure N.N: description — filename.png]*  (non-screenshot)
        m_img5 = re.match(r'\*\[Figure (\d+\.\d+): ([^—\]]+) — ([^\]]+\.(?:png|jpg|jpeg))\]\*', stripped)
        if m_img5:
            num = m_img5.group(1); desc = m_img5.group(2).strip(); fname = m_img5.group(3).strip()
            if num in SHOT_FIGURES:
                path, cap = SHOT_FIGURES[num]
                add_img(doc, path, cap, w=14.0)
            else:
                path = os.path.join(IMGS, fname)
                add_img(doc, path, f"Figure {num}: {desc}", w=13.0)
            i+=1; continue

        # headings
        if stripped.startswith('#### '): sec_head(doc, stripped[5:], level=3); i+=1; continue
        if stripped.startswith('### '):
            txt = stripped[4:]
            if re.match(r'^5\.3\b', txt):
                in_53=True; sec_head(doc, txt, level=1)
                for f in FINDINGS: finding_card(doc, f)
                i+=1
                while i < len(lines):
                    l2=lines[i].strip()
                    if re.match(r'^#{2,3}\s+5\.[4-9]',l2) or re.match(r'^#{1,3}\s+6\.',l2):
                        in_53=False; break
                    i+=1
                continue
            sec_head(doc, txt, level=2); i+=1; continue

        if stripped.startswith('## '):
            txt = stripped[3:]; in_53=False
            # detect appendix sections
            m_app = re.match(r'^Appendix ([A-E]):', txt)
            if m_app:
                in_appendix=True; current_appendix=m_app.group(1)
                sec_head(doc, txt, level=1)
                # insert appendix screenshots if we have them
                shots = APPENDIX_SHOTS.get(current_appendix, [])
                for path, cap in shots:
                    add_img(doc, path, cap, w=14.0)
                i+=1; continue
            sec_head(doc, txt, level=1); i+=1; continue

        if stripped.startswith('# '): ch_head(doc, stripped[2:]); i+=1; continue

        # bullet list  (- item  or  * item)
        m_b = re.match(r'^[-*]\s+(.+)', stripped)
        if m_b and not in_53:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent=Cm(1.27); p.paragraph_format.first_line_indent=Cm(0)
            p.paragraph_format.space_after = Pt(3)
            add_inline_md(p, m_b.group(1).strip())
            i+=1; continue

        # numbered list — render as body paragraph with number prefix (same font as body)
        m_num = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if m_num and not in_53:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(4)
            num_run = p.add_run(f"{m_num.group(1)}.  ")
            fmt(num_run, size=12)
            add_inline_md(p, m_num.group(2).strip())
            i+=1; continue

        # skip front-matter meta lines and blockquotes
        if any(stripped.startswith(x) for x in ['**Student','**Submission','**Program',
           '**Institution','**Supervisor','> ','## Declaration','## Abstract','## Table',
           '**Address','**Website','**Email','**Telephone','**Keywords']):
            i+=1; continue

        if not stripped: i+=1; continue
        if in_53: i+=1; continue

        # normal paragraph — render with inline markdown formatting
        if stripped.startswith('*[') or stripped.startswith('['):
            # unmatched figure/table reference — skip silently
            i+=1; continue
        body(doc, stripped)
        i+=1

# ── Build ─────────────────────────────────────────────────────────────────────
def build():
    # Report missing files upfront
    missing = []
    for key,(path,cap) in SHOT_FIGURES.items():
        if path is None or not os.path.exists(path):
            bname = os.path.basename(path) if path else ("fig" + key.replace(".", "") + "_*.png")
            missing.append("  screenshots/" + bname + "  ->  " + cap)
    for key,(path,cap) in LOGO_FIGURES.items():
        if not os.path.exists(path):
            missing.append(f"  image/{os.path.basename(path)}  →  {cap}")
    if missing:
        print("⚠  MISSING FILES (shown as grey placeholder in doc):")
        for m in missing: print(m)
        print()

    doc = new_doc()
    add_front_matter(doc)
    process_md(doc)
    doc.save(OUT)
    print(f"✓  Saved → {OUT}")
    print("   Open in Word → Ctrl+A → F9 to update Table of Contents")

if __name__ == '__main__':
    build()
